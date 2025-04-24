# import os
# import re
# from flask import Flask, request, jsonify
# from flask_cors import CORS
# from transformers import pipeline
# from docx import Document
# import fitz  # PyMuPDF
# from dotenv import load_dotenv
# from together import Together
# from googletrans import Translator  # Import the translator

# # Load environment variables
# load_dotenv()

# # Initialize Flask app
# app = Flask(__name__)
# CORS(app)  # 🔥 Enables CORS for all origins

# # Load emotion classifier
# emotion_classifier = pipeline(
#     "text-classification",
#     model="j-hartmann/emotion-english-distilroberta-base",
#     top_k=1
# )

# # Map emotions to emojis
# emotion_to_emoji = {
#     "joy": "😊", "anger": "😠", "sadness": "😢",
#     "fear": "😨", "surprise": "😲", "disgust": "🤢",
#     "neutral": "💬", "confusion": "🤔", "love": "❤️",
# }

# # Initialize the Google Translator
# translator = Translator()

# # Extract text from uploaded files
# def extract_text(file_path):
#     if file_path.endswith(".pdf"):
#         doc = fitz.open(file_path)
#         return "\n".join([page.get_text() for page in doc])
#     elif file_path.endswith(".docx"):
#         doc = Document(file_path)
#         return "\n".join([p.text for p in doc.paragraphs])
#     elif file_path.endswith(".txt"):
#         with open(file_path, "r", encoding="utf-8") as f:
#             return f.read()
#     else:
#         raise ValueError("Unsupported file format")

# # Clean the response
# def clean_response(answer):
#     # Remove anything starting with Note:, Translation:, or within parentheses mentioning translation
#     answer = re.split(r'\b(Note|Translation)\b:', answer, flags=re.IGNORECASE)[0]
#     answer = re.sub(r'\(Note:.*?\)', '', answer, flags=re.IGNORECASE)
#     return answer.strip()

# # Translate text to any language
# def translate_text(text, target_language):
#     # If the target language is the same as the input, no translation is necessary.
#     if target_language == 'en':
#         return text

#     # Translate text
#     translation = translator.translate(text, dest=target_language)
#     print(f"Translated text from English to {target_language}: {translation.text}")  # Debugging line
#     return translation.text

# # Ask question using Together API
# def ask_question(content, question, target_language):
#     prompt = f"""
# You are an intelligent assistant helping the user understand an uploaded document.

# The user may ask questions related to the document, and you should answer based on the content provided. Avoid adding any translation disclaimers, notes, or similar metadata in your response. Just give the direct answer.

# --- BEGIN CONTENT ---
# {content}
# --- END CONTENT ---

# User Question: {question}
# """
#     client = Together(api_key=os.getenv("TOGETHER_API_KEY"))
#     response = client.chat.completions.create(
#         model="meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
#         messages=[{"role": "user", "content": prompt}]
#     )
#     raw_answer = response.choices[0].message.content.strip()

#     # Emotion detection
#     try:
#         emotion_result = emotion_classifier(raw_answer)
#         emotion = emotion_result[0]['label'].lower() if isinstance(emotion_result, list) else emotion_result['label'].lower()
#         emoji = emotion_to_emoji.get(emotion, "")
#     except Exception as e:
#         print(f"Emotion classification error: {e}")
#         emoji = ""

#     cleaned_answer = clean_response(raw_answer)

#     # Translate the answer to the target language
#     translated_answer = translate_text(cleaned_answer, target_language)
#     return translated_answer

# # Route: Accepts a file and a question
# @app.route("/ask", methods=["POST"])
# def ask_with_file():
#     if "file" not in request.files or "question" not in request.form:
#         return jsonify({"error": "File and question are required"}), 400

#     file = request.files["file"]
#     question = request.form["question"]
#     target_language = request.form.get("language", "en")  # Default to English if no language is provided
#     filepath = os.path.join("uploads", file.filename)

#     os.makedirs("uploads", exist_ok=True)
#     file.save(filepath)

#     print(f"User specified target language: {target_language}")  # Debugging line

#     try:
#         content = extract_text(filepath)
#         answer = ask_question(content, question, target_language)
#         return jsonify({"answer": answer})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# # Route: Uses predefined file
# @app.route("/ask_predefined", methods=["POST"])
# def ask_with_predefined_file():
#     question = request.form.get("question")
#     target_language = request.form.get("language", "en")  # Default to English if no language is provided
#     if not question:
#         return jsonify({"error": "Question is required"}), 400

#     predefined_file_path = os.path.join(os.getcwd(), "uploads", "Sample_BGV_Document.pdf")

#     if not os.path.exists(predefined_file_path):
#         return jsonify({"error": "Predefined file not found in the uploads directory"}), 404

#     print(f"User specified target language: {target_language}")  # Debugging line

#     try:
#         content = extract_text(predefined_file_path)
#         answer = ask_question(content, question, target_language)
#         return jsonify({"answer": answer})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# # Root route
# @app.route("/", methods=["GET"])
# def home():
#     return "<h2>📄 Document Q&A API is Running</h2>"

# # Start app
# if __name__ == "__main__":
#     port = int(os.getenv("PORT", 5000))
#     app.run(host="0.0.0.0", port=port) 









import os
import re
from flask import Flask, request, jsonify
from flask_cors import CORS
from transformers import pipeline, AutoModelForSequenceClassification, AutoTokenizer
from docx import Document
import fitz  # PyMuPDF
from dotenv import load_dotenv
from together import Together
from googletrans import Translator  # Import the translator

# Load environment variables
load_dotenv()

# Initialize Flask app
app = Flask(__name__)
CORS(app)  # 🔥 Enables CORS for all origins

# Load emotion classifier for text
# emotion_classifier = pipeline(
#     "text-classification",
#     model="j-hartmann/emotion-english-distilroberta-base",
#     top_k=1
# )

# # Map emotions to emojis
# emotion_to_emoji = {
#     "joy": "😊", "anger": "😠", "sadness": "😢",
#     "fear": "😨", "surprise": "😲", "disgust": "🤢",
#     "neutral": "💬", "confusion": "🤔", "love": "❤️",
# }

# Initialize the Google Translator
translator = Translator()

# Load the KoalaAI Emoji Suggestion Model
# emoji_model = AutoModelForSequenceClassification.from_pretrained("KoalaAI/Emoji-Suggester")
# emoji_tokenizer = AutoTokenizer.from_pretrained("KoalaAI/Emoji-Suggester")

# # Create the emoji suggestion pipeline
# emoji_pipe = pipeline("text-classification", model=emoji_model, tokenizer=emoji_tokenizer)

# Extract text from uploaded files
def extract_text(file_path):
    if file_path.endswith(".pdf"):
        doc = fitz.open(file_path)
        return "\n".join([page.get_text() for page in doc])
    elif file_path.endswith(".docx"):
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file_path.endswith(".txt"):
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError("Unsupported file format")

# Clean the response
def clean_response(answer):
    # Remove anything starting with Note:, Translation:, or within parentheses mentioning translation
    answer = re.split(r'\b(Note|Translation)\b:', answer, flags=re.IGNORECASE)[0]
    answer = re.sub(r'\(Note:.*?\)', '', answer, flags=re.IGNORECASE)
    return answer.strip()

# Translate text to any language
def translate_text(text, target_language):
    if target_language == 'en':
        return text
    translation = translator.translate(text, dest=target_language)
    print(f"Translated text from English to {target_language}: {translation.text}")  # Debugging line
    return translation.text

# Ask question using Together API
def ask_question(content, question, target_language):
    prompt = f"""
You are an intelligent assistant helping the user understand an uploaded document.

The user may ask questions related to the document, and you should answer based on the content provided. Avoid adding any translation disclaimers, notes, or similar metadata in your response. Just give the direct answer.

--- BEGIN CONTENT ---
{content}
--- END CONTENT ---

User Question: {question}
"""
    client = Together(api_key=os.getenv("TOGETHER_API_KEY"))
    response = client.chat.completions.create(
        model="meta-llama/Llama-3.3-70B-Instruct-Turbo-Free",
        messages=[{"role": "user", "content": prompt}]
    )
    raw_answer = response.choices[0].message.content.strip()

    # Emotion detection
    # try:
    #     emotion_result = emotion_classifier(raw_answer)
    #     emotion = emotion_result[0]['label'].lower() if isinstance(emotion_result, list) else emotion_result['label'].lower()
    #     emoji = emotion_to_emoji.get(emotion, "")
    # except Exception as e:
    #     print(f"Emotion classification error: {e}")
    #     emoji = ""

    cleaned_answer = clean_response(raw_answer)

    # Translate the answer to the target language
    translated_answer = translate_text(cleaned_answer, target_language)
    
    # Suggest emojis based on the answer
    # emoji_suggestion = emoji_pipe(translated_answer)
    # emojis = [emoji['label'] for emoji in emoji_suggestion]

    # Append emojis to the answer
    final_answer = f"{translated_answer} "
    return final_answer

# Route: Accepts a file and a question
@app.route("/ask", methods=["POST"])
def ask_with_file():
    if "file" not in request.files or "question" not in request.form:
        return jsonify({"error": "File and question are required"}), 400

    file = request.files["file"]
    question = request.form["question"]
    target_language = request.form.get("language", "en")  # Default to English if no language is provided
    filepath = os.path.join("uploads", file.filename)

    os.makedirs("uploads", exist_ok=True)
    file.save(filepath)

    print(f"User specified target language: {target_language}")  # Debugging line

    try:
        content = extract_text(filepath)
        answer = ask_question(content, question, target_language)
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Route: Uses predefined file
@app.route("/ask_predefined", methods=["POST"])
def ask_with_predefined_file():
    question = request.form.get("question")
    target_language = request.form.get("language", "en")  # Default to English if no language is provided
    if not question:
        return jsonify({"error": "Question is required"}), 400

    predefined_file_path = os.path.join(os.getcwd(), "uploads", "Sample_BGV_Document.pdf")

    if not os.path.exists(predefined_file_path):
        return jsonify({"error": "Predefined file not found in the uploads directory"}), 404

    print(f"User specified target language: {target_language}")  # Debugging line

    try:
        content = extract_text(predefined_file_path)
        answer = ask_question(content, question, target_language)
        return jsonify({"answer": answer})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Root route
@app.route("/", methods=["GET"])
def home():
    return "<h2>📄 Document Q&A API is Running</h2>"

# Start app
if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
